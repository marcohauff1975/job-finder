"""
Job Search - Job Finder + Company Researcher + Resume Tailor agents.

Job Finder searches the live web (via Serper) for job postings matching
your criteria. Company Researcher takes a company name from those
results and researches it. Resume Tailor takes a resume plus a specific
job posting and that research, and produces a tailored resume.

This file is written to support multiple users: every function that
touches a resume or search history takes an explicit path, rather than
using one fixed file for everybody. streamlit_app.py figures out each
logged-in user's own paths (under users/<username>/) and passes them
in.

The agents and tasks themselves are defined in plain text in:
    config/agents.yaml
    config/tasks.yaml
This file just loads those definitions, wires them together, and adds
the supporting logic (new-posting tracking, resume reading/writing).

Run directly for a quick terminal test (uses the DEMO_* constants
below, not any particular logged-in user):
    python job_search.py
"""

import io
import json
import os
import re
from pathlib import Path

import yaml
from crewai import Agent, Task, Crew, Process, LLM
from crewai_tools import SerperDevTool
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel
from pypdf import PdfReader

load_dotenv()


# --- Structured output shapes -----------------------------------------
# Defining these tells CrewAI exactly what fields each result must
# have, so results are reliable data (not just paragraphs of text).

class JobPosting(BaseModel):
    title: str
    company: str
    location: str
    link: str
    salary: str = ""


class JobSearchResult(BaseModel):
    postings: list[JobPosting]


class CompanyResearch(BaseModel):
    overview: str
    size: str = ""
    funding: str = ""
    reputation: str = ""
    other_open_roles: str = ""
    tech_stack: str = ""
    recent_news: str = ""


class TailoredResume(BaseModel):
    tailored_paragraphs: list[str]
    changes_summary: str


class ResumeReview(BaseModel):
    strengths: list[str]
    weaknesses: list[str]
    suggestions: list[str]


class ExperienceEntry(BaseModel):
    title: str
    company: str
    location: str = ""
    dates: str = ""
    bullets: list[str] = []


class EducationEntry(BaseModel):
    degree: str
    school: str
    dates: str = ""


class ResumeContent(BaseModel):
    """A resume broken into clean structured data, used to rebuild it in
    a different visual template. Extracted once by the resume_formatter
    agent, then handed to any of the plain-Python template renderers
    below - the LLM only extracts data, it never lays out documents."""
    name: str
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    summary: str = ""
    experience: list[ExperienceEntry] = []
    education: list[EducationEntry] = []
    skills: list[str] = []


# --- Demo search criteria -------------------------------------------------
# Only used when running this file directly (python job_search.py) for a
# quick terminal test. The Streamlit UI always lets the logged-in user
# set these instead.

DEMO_ROLE = "CTO"
DEMO_LOCATION = "Amsterdam, Netherlands"
DEMO_REMOTE = True
DEMO_RESUME_PATH = Path(__file__).parent / "resume" / "CV_Marco_Hauff_2026_v6.docx"
DEMO_HISTORY_DIR = Path(__file__).parent / "history" / "demo"

# --- Load agent/task definitions from the config/ folder -----------------

CONFIG_DIR = Path(__file__).parent / "config"

with open(CONFIG_DIR / "agents.yaml", "r") as f:
    agents_config = yaml.safe_load(f)

with open(CONFIG_DIR / "tasks.yaml", "r") as f:
    tasks_config = yaml.safe_load(f)

# --- LLM (same Claude setup as main.py) --------------------------------

claude = LLM(model="anthropic/claude-sonnet-5")

# --- Tool ---------------------------------------------------------------
# SerperDevTool lets agents run real Google searches.
# It reads SERPER_API_KEY from the environment automatically.

search_tool = SerperDevTool()

# --- Job Finder: agent + task + crew ---------------------------------------

job_finder = Agent(
    config=agents_config["job_finder"],
    tools=[search_tool],
    llm=claude,
    verbose=True,
    inject_date=True,  # gives the agent today's date, so it can judge
                        # what counts as "posted in the last two weeks"
    max_iter=8,  # safety net on top of the task's own search budget -
                 # without this the agent kept digging into individual
                 # postings (salary, exact date, full description) for
                 # minutes; CrewAI's default of 20 was way too loose.
)

job_search_task = Task(
    config=tasks_config["job_search_task"],
    agent=job_finder,
    output_pydantic=JobSearchResult,
)

job_search_crew = Crew(
    agents=[job_finder],
    tasks=[job_search_task],
    process=Process.sequential,
    verbose=True,
)

# --- Company Researcher: agent + task + crew --------------------------------

company_researcher = Agent(
    config=agents_config["company_researcher"],
    tools=[search_tool],
    llm=claude,
    verbose=True,
)

company_research_task = Task(
    config=tasks_config["company_research_task"],
    agent=company_researcher,
    output_pydantic=CompanyResearch,
)

research_crew = Crew(
    agents=[company_researcher],
    tasks=[company_research_task],
    process=Process.sequential,
    verbose=True,
)

# --- Resume Tailor: agent + task + crew --------------------------------------
# No search tool needed here - it only reasons over the resume, job
# posting, and company research it's given.

resume_tailor = Agent(
    config=agents_config["resume_tailor"],
    llm=claude,
    verbose=True,
)

resume_tailor_task = Task(
    config=tasks_config["resume_tailor_task"],
    agent=resume_tailor,
    output_pydantic=TailoredResume,
)

resume_crew = Crew(
    agents=[resume_tailor],
    tasks=[resume_tailor_task],
    process=Process.sequential,
    verbose=True,
)

# --- Resume Reviewer: agent + task + crew ------------------------------
# General-purpose feedback on the resume as it stands - not tied to any
# specific job posting, unlike the tailor above.

resume_reviewer = Agent(
    config=agents_config["resume_reviewer"],
    llm=claude,
    verbose=True,
)

resume_review_task = Task(
    config=tasks_config["resume_review_task"],
    agent=resume_reviewer,
    output_pydantic=ResumeReview,
)

review_crew = Crew(
    agents=[resume_reviewer],
    tasks=[resume_review_task],
    process=Process.sequential,
    verbose=True,
)

# --- Resume Formatter: agent + task + crew -----------------------------
# This agent only extracts the resume into structured data - it never
# lays out a document. The actual visual templates (below, under
# "Format renderers") are plain Python/python-docx, so the layout is
# deterministic and never varies between runs.

resume_formatter = Agent(
    config=agents_config["resume_formatter"],
    llm=claude,
    verbose=True,
)

resume_format_extract_task = Task(
    config=tasks_config["resume_format_extract_task"],
    agent=resume_formatter,
    output_pydantic=ResumeContent,
)

format_extract_crew = Crew(
    agents=[resume_formatter],
    tasks=[resume_format_extract_task],
    process=Process.sequential,
    verbose=True,
)

# --- Shared helpers -----------------------------------------------------

def _slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")


# --- "New since last search" tracking -----------------------------------
# Each role+location combo gets its own small history file within the
# given history_dir (streamlit_app.py passes a per-user folder), so
# running a search for "Product Manager" in "Berlin" is tracked
# separately from "Head of Technology" in "Munich" - and separately per
# user, since each user has their own history_dir.

def _history_file(history_dir: Path, role: str, location: str) -> Path:
    history_dir.mkdir(parents=True, exist_ok=True)
    return history_dir / f"{_slugify(f'{role}_{location}')}.json"


def _posting_key(job: JobPosting) -> str:
    # Prefer the link as the unique identifier; fall back to title+company
    # if a posting somehow has no link.
    return job.link or f"{job.title}|{job.company}"


# --- Resume reading -------------------------------------------------------
# Every function below takes an explicit resume_path, since each user
# has their own uploaded resume file (see streamlit_app.py).
#
# Instead of throwing away a resume's design and building a plain new
# document, tailoring works by editing the TEXT of the original file's
# existing paragraphs in place - so fonts, bullet styles, and any
# header/skills tables stay exactly as they were. That only works
# reliably if the tailored content lines up 1:1 with the original
# paragraphs (same order, same count), which is why the task instructs
# the agent not to add, remove, or reorder paragraphs - only reword them.

def _pdf_to_docx_bytes(pdf_bytes: bytes) -> bytes:
    """Converts a PDF's text into a plain .docx (one paragraph per
    non-blank line). This is deliberately simple rather than
    layout-preserving: everything downstream (tailoring, format
    rendering, extraction) only ever reads a resume as a list of
    paragraphs, so a clean paragraph-per-line .docx is exactly what
    those steps need, regardless of the PDF's original visual
    layout."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    doc = DocxDocument()
    for line in text.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def save_resume_upload(filename: str, file_bytes: bytes, resume_path: Path) -> None:
    """Saves an uploaded resume as resume.docx, converting it first if
    it isn't already a .docx (currently: .pdf). Every other function in
    this file assumes resume_path is a real .docx it can open with
    python-docx, so conversion has to happen exactly once, here, at
    upload time."""
    if filename.lower().endswith(".pdf"):
        resume_path.write_bytes(_pdf_to_docx_bytes(file_bytes))
    else:
        resume_path.write_bytes(file_bytes)


def _load_resume_doc(resume_path: Path) -> DocxDocument:
    if not resume_path.exists():
        raise FileNotFoundError(f"No resume found at {resume_path}.")
    return DocxDocument(resume_path)


def _editable_paragraphs(doc: DocxDocument) -> list:
    """The paragraphs that make up the resume's flowing text (skips blank
    spacer paragraphs, and skips table content like header/skills boxes,
    which aren't tailored)."""
    return [p for p in doc.paragraphs if p.text.strip()]


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph's text while keeping its existing formatting
    (font, bold, size, bullet style, etc.) - by writing into its first
    run rather than creating a new plain one."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def read_resume_paragraphs(resume_path: Path) -> list[str]:
    """The resume's tailorable paragraphs, as plain text, in order."""
    doc = _load_resume_doc(resume_path)
    return [p.text for p in _editable_paragraphs(doc)]


def build_tailored_docx_bytes(
    resume_path: Path, tailored_paragraphs: list[str]
) -> bytes:
    """Build the tailored resume as .docx bytes (built from the ORIGINAL
    file, with only the text of each paragraph swapped in, so formatting
    is preserved) - ready to hand to Streamlit's download_button. Nothing
    is written to disk here; the caller decides what to do with the
    bytes."""
    doc = _load_resume_doc(resume_path)
    editable = _editable_paragraphs(doc)
    # zip() stops at the shorter list, so if the agent returned a
    # different count than expected, the extra paragraphs (either side)
    # are simply left as-is rather than crashing.
    for paragraph, new_text in zip(editable, tailored_paragraphs):
        _set_paragraph_text(paragraph, new_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --- Formatting helpers for feeding structured data into task prompts ----

def _format_job_posting(job: dict) -> str:
    return (
        f"Title: {job['title']}\n"
        f"Company: {job['company']}\n"
        f"Location: {job['location']}\n"
        f"Salary: {job['salary'] or 'not specified'}\n"
        f"Link: {job['link']}"
    )


def _format_company_research(research: CompanyResearch) -> str:
    lines = [f"Overview: {research.overview}"]
    if research.size:
        lines.append(f"Size: {research.size}")
    if research.funding:
        lines.append(f"Funding: {research.funding}")
    if research.reputation:
        lines.append(f"Reputation: {research.reputation}")
    if research.other_open_roles:
        lines.append(f"Other open roles: {research.other_open_roles}")
    if research.tech_stack:
        lines.append(f"Tech stack: {research.tech_stack}")
    if research.recent_news:
        lines.append(f"Recent news: {research.recent_news}")
    return "\n".join(lines)


# --- Reusable functions --------------------------------------------------
# The terminal script below and the Streamlit UI (streamlit_app.py) call
# these same functions.

def find_jobs(
    role: str, location: str, remote: bool, history_dir: Path
) -> list[dict]:
    """Run the job search crew and return a list of posting dicts, each
    marked with is_new (True if not seen in the last search for this
    same role+location, within this history_dir)."""
    remote_note = (
        "The person is open to fully remote roles."
        if remote
        else "The person wants on-site or hybrid roles only."
    )
    inputs = {
        "role": role,
        "location": location,
        "remote_note": remote_note,
    }
    result = job_search_crew.kickoff(inputs=inputs)
    postings = result.pydantic.postings if result.pydantic else []

    history_path = _history_file(history_dir, role, location)
    previous_keys = set()
    if history_path.exists():
        with open(history_path, "r") as f:
            previous_keys = set(json.load(f))

    enriched = []
    current_keys = []
    for job in postings:
        key = _posting_key(job)
        current_keys.append(key)
        enriched.append({
            "title": job.title,
            "company": job.company,
            "location": job.location,
            "link": job.link,
            "salary": job.salary,
            "is_new": key not in previous_keys,
        })

    with open(history_path, "w") as f:
        json.dump(current_keys, f)

    return enriched


def research_company(company: str, role: str) -> CompanyResearch | None:
    """Run the company research crew for a single company and return the
    structured result (or None if the agent didn't return valid data)."""
    inputs = {"company": company, "role": role}
    result = research_crew.kickoff(inputs=inputs)
    return result.pydantic if result.pydantic else None


def tailor_resume_for_job(
    job: dict, company_research: CompanyResearch, resume_path: Path
) -> TailoredResume | None:
    """Run the resume tailor crew for a specific job posting + company
    research, using the resume at resume_path, and return the structured
    result (or None if it failed)."""
    paragraphs = read_resume_paragraphs(resume_path)
    numbered_resume = "\n".join(
        f"{i + 1}. {text}" for i, text in enumerate(paragraphs)
    )

    inputs = {
        "resume": numbered_resume,
        "paragraph_count": str(len(paragraphs)),
        "role": job["title"],
        "company": job["company"],
        "job_posting": _format_job_posting(job),
        "company_research": _format_company_research(company_research),
    }
    result = resume_crew.kickoff(inputs=inputs)
    return result.pydantic if result.pydantic else None


def review_resume(resume_path: Path) -> ResumeReview | None:
    """Run the resume review crew (general feedback, not tied to any
    specific job) and return the structured result, or None if it
    failed."""
    paragraphs = read_resume_paragraphs(resume_path)
    inputs = {"resume_text": "\n".join(paragraphs)}
    result = review_crew.kickoff(inputs=inputs)
    return result.pydantic if result.pydantic else None


def extract_resume_content(resume_path: Path) -> ResumeContent | None:
    """Run the resume formatter crew to extract structured resume data
    (name, contact info, summary, experience, education, skills), ready
    to be rebuilt in a different visual template. Returns None if
    extraction failed."""
    paragraphs = read_resume_paragraphs(resume_path)
    inputs = {"resume_text": "\n".join(paragraphs)}
    result = format_extract_crew.kickoff(inputs=inputs)
    return result.pydantic if result.pydantic else None


# --- Format renderers -----------------------------------------------------
# Plain python-docx code - deliberately NOT an LLM step. resume_formatter
# (above) extracts the content once; these functions lay it out. Keeping
# layout as ordinary code (not agent output) means the same content
# always renders identically, with no risk of the model inventing or
# dropping something during layout.

def _add_run(paragraph, text, *, size=10, bold=False, italic=False,
             color=None, font_name="Calibri"):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _set_cell_background(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _remove_cell_borders(cell) -> None:
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "FFFFFF")
        borders.append(el)
    cell._tc.get_or_add_tcPr().append(borders)


def _add_bottom_border(paragraph, color: str = "D3D1C7", size: int = 4) -> None:
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pBdr)


def _add_right_tab(paragraph, position_in: float = 7.0) -> None:
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(position_in), WD_TAB_ALIGNMENT.RIGHT
    )


def render_executive_signature_docx(content: ResumeContent) -> bytes:
    """Layout based directly on Marco's own resume design: a bold name +
    title header, navy ALL-CAPS section labels, a 3-column
    core-competencies grid for skills, and company/date lines with a
    right-aligned tab. No photo - this app doesn't collect one, and it's
    being offered as a template to users other than Marco."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.56)
    section.bottom_margin = Inches(0.56)

    NAVY, BODY, GRAY = "1F4E79", "1A1A1A", "666666"

    p = doc.add_paragraph()
    _add_run(p, content.name or "Your Name", size=20, bold=True, color=BODY)
    current_title = content.experience[0].title if content.experience else ""
    if current_title:
        p = doc.add_paragraph()
        _add_run(p, current_title, size=12, bold=True, color=NAVY)

    contact_bits = [b for b in (content.location, content.phone, content.email) if b]
    if contact_bits:
        p = doc.add_paragraph()
        _add_run(p, "  ·  ".join(contact_bits), size=10, color=GRAY)
    if content.linkedin:
        p = doc.add_paragraph()
        _add_run(p, content.linkedin, size=10, color=GRAY)

    def section_header(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, text.upper(), size=11, bold=True, color=NAVY)
        return p

    if content.summary:
        section_header("Profile")
        p = doc.add_paragraph()
        _add_run(p, content.summary, size=10, color=BODY)

    if content.skills:
        section_header("Core Competencies")
        rows = [content.skills[i:i + 3] for i in range(0, len(content.skills), 3)]
        table = doc.add_table(rows=len(rows), cols=3)
        for r, row_skills in enumerate(rows):
            for c in range(3):
                cell = table.rows[r].cells[c]
                _remove_cell_borders(cell)
                if c < len(row_skills):
                    _add_run(cell.paragraphs[0], f"▸  {row_skills[c]}", size=10, color=BODY)

    if content.experience:
        section_header("Work Experience")
        for job in content.experience:
            p = doc.add_paragraph()
            _add_right_tab(p)
            company_line = job.company + (f"  ({job.location})" if job.location else "")
            _add_run(p, company_line, size=10, bold=True, color=BODY)
            if job.dates:
                _add_run(p, f"\t{job.dates}", size=10, color=GRAY)
            if job.title:
                p = doc.add_paragraph()
                _add_run(p, job.title, size=10, color=GRAY)
            for bullet in job.bullets:
                bp = doc.add_paragraph(style="List Bullet")
                _add_run(bp, bullet, size=10, color=BODY)

    if content.education:
        section_header("Education")
        for edu in content.education:
            p = doc.add_paragraph()
            _add_right_tab(p)
            text = edu.degree + (f" — {edu.school}" if edu.school else "")
            _add_run(p, text, size=10, bold=True, color=BODY)
            if edu.dates:
                _add_run(p, f"\t{edu.dates}", size=10, color=GRAY)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_modern_minimalist_docx(content: ResumeContent) -> bytes:
    """Clean single-column layout: bold name, muted contact line, a soft
    purple accent on section headers with a thin rule underneath."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    ACCENT, BODY, MUTED = "534AB7", "1A1A1A", "5F5E5A"

    p = doc.add_paragraph()
    _add_run(p, content.name or "Your Name", size=24, bold=True, color=BODY)

    contact_bits = [b for b in (content.location, content.phone, content.email, content.linkedin) if b]
    if contact_bits:
        p = doc.add_paragraph()
        _add_run(p, "   ".join(contact_bits), size=10, color=MUTED)

    def section_header(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, text.upper(), size=11, bold=True, color=ACCENT)
        _add_bottom_border(p)
        return p

    if content.summary:
        section_header("Summary")
        p = doc.add_paragraph()
        _add_run(p, content.summary, size=10.5, color=BODY)

    if content.experience:
        section_header("Experience")
        for job in content.experience:
            p = doc.add_paragraph()
            _add_run(p, job.title or "", size=11, bold=True, color=BODY)
            meta_bits = [b for b in (job.company, job.location, job.dates) if b]
            if meta_bits:
                p = doc.add_paragraph()
                _add_run(p, "  |  ".join(meta_bits), size=10, italic=True, color=MUTED)
            for bullet in job.bullets:
                bp = doc.add_paragraph(style="List Bullet")
                _add_run(bp, bullet, size=10, color=BODY)

    if content.education:
        section_header("Education")
        for edu in content.education:
            p = doc.add_paragraph()
            text = edu.degree + (f", {edu.school}" if edu.school else "")
            text += f"  ({edu.dates})" if edu.dates else ""
            _add_run(p, text, size=10.5, color=BODY)

    if content.skills:
        section_header("Skills")
        p = doc.add_paragraph()
        _add_run(p, "  •  ".join(content.skills), size=10, color=BODY)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_two_column_docx(content: ResumeContent) -> bytes:
    """Shaded sidebar (contact, skills, education) alongside a main
    column (headline, summary, experience), built with a borderless
    two-column table."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    SIDEBAR_BG, ACCENT, BODY, MUTED = "EEF1F4", "185FA5", "1A1A1A", "5F5E5A"

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.3)
    table.columns[1].width = Inches(4.7)
    sidebar, main = table.rows[0].cells
    sidebar.width = Inches(2.3)
    main.width = Inches(4.7)
    _remove_cell_borders(sidebar)
    _remove_cell_borders(main)
    _set_cell_background(sidebar, SIDEBAR_BG)

    _add_run(sidebar.paragraphs[0], content.name or "Your Name", size=15, bold=True, color=BODY)

    def sidebar_header(text: str):
        p = sidebar.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        _add_run(p, text.upper(), size=10, bold=True, color=ACCENT)

    sidebar_header("Contact")
    for bit in (content.location, content.phone, content.email, content.linkedin):
        if bit:
            _add_run(sidebar.add_paragraph(), bit, size=9, color=MUTED)

    if content.skills:
        sidebar_header("Skills")
        for skill in content.skills:
            _add_run(sidebar.add_paragraph(), skill, size=9, color=BODY)

    if content.education:
        sidebar_header("Education")
        for edu in content.education:
            text = edu.degree + (f", {edu.school}" if edu.school else "")
            _add_run(sidebar.add_paragraph(), text, size=9, bold=True, color=BODY)
            if edu.dates:
                _add_run(sidebar.add_paragraph(), edu.dates, size=9, color=MUTED)

    current_title = content.experience[0].title if content.experience else ""
    if current_title:
        _add_run(main.paragraphs[0], current_title, size=13, bold=True, color=ACCENT)

    def main_header(text: str):
        p = main.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        _add_run(p, text.upper(), size=11, bold=True, color=BODY)
        _add_bottom_border(p)

    if content.summary:
        main_header("Summary")
        _add_run(main.add_paragraph(), content.summary, size=10, color=BODY)

    if content.experience:
        main_header("Experience")
        for job in content.experience:
            _add_run(main.add_paragraph(), job.title or "", size=10.5, bold=True, color=BODY)
            meta_bits = [b for b in (job.company, job.dates) if b]
            if meta_bits:
                _add_run(main.add_paragraph(), "  ·  ".join(meta_bits), size=9.5, italic=True, color=MUTED)
            for bullet in job.bullets:
                bp = main.add_paragraph(style="List Bullet")
                _add_run(bp, bullet, size=9.5, color=BODY)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_classic_executive_docx(content: ResumeContent) -> bytes:
    """Traditional, centered, serif layout with horizontal rules - the
    formal/conservative option for industries where the other templates
    might read as too casual."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    BODY, MUTED, FONT = "1A1A1A", "444441", "Georgia"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, (content.name or "Your Name").upper(), size=18, bold=True, color=BODY, font_name=FONT)

    contact_bits = [b for b in (content.location, content.phone, content.email, content.linkedin) if b]
    if contact_bits:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, "  |  ".join(contact_bits), size=10, color=MUTED, font_name=FONT)
        _add_bottom_border(p, color="2C2C2A", size=6)

    def section_header(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, text.upper(), size=11, bold=True, color=BODY, font_name=FONT)

    if content.summary:
        section_header("Professional Summary")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_run(p, content.summary, size=10.5, color=BODY, font_name=FONT)

    if content.experience:
        section_header("Professional Experience")
        for job in content.experience:
            p = doc.add_paragraph()
            _add_run(p, job.title or "", size=11, bold=True, color=BODY, font_name=FONT)
            meta_line = ", ".join(b for b in (job.company, job.location) if b)
            if job.dates:
                meta_line = f"{meta_line}  ({job.dates})" if meta_line else job.dates
            if meta_line:
                p = doc.add_paragraph()
                _add_run(p, meta_line, size=10, italic=True, color=MUTED, font_name=FONT)
            for bullet in job.bullets:
                bp = doc.add_paragraph(style="List Bullet")
                _add_run(bp, bullet, size=10, color=BODY, font_name=FONT)

    if content.education:
        section_header("Education")
        for edu in content.education:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = edu.degree + (f", {edu.school}" if edu.school else "")
            text += f"  ({edu.dates})" if edu.dates else ""
            _add_run(p, text, size=10.5, color=BODY, font_name=FONT)

    if content.skills:
        section_header("Core Competencies")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, "  •  ".join(content.skills), size=10, color=BODY, font_name=FONT)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# Registry the UI uses to list format choices and dispatch to the right
# renderer, without the UI needing to know about python-docx internals.
FORMAT_TEMPLATES = {
    "executive_signature": {
        "label": "Executive signature",
        "description": "Based on Marco's own resume - navy section headers and a core-competencies grid.",
        "render": render_executive_signature_docx,
    },
    "modern_minimalist": {
        "label": "Modern minimalist",
        "description": "Clean single-column layout with a soft purple accent.",
        "render": render_modern_minimalist_docx,
    },
    "two_column": {
        "label": "Two-column",
        "description": "Shaded sidebar for contact, skills, and education; main column for experience.",
        "render": render_two_column_docx,
    },
    "classic_executive": {
        "label": "Classic executive",
        "description": "Traditional centered layout, serif font, formal tone.",
        "render": render_classic_executive_docx,
    },
}


def render_resume_in_format(content: ResumeContent, template_key: str) -> bytes:
    return FORMAT_TEMPLATES[template_key]["render"](content)


if __name__ == "__main__":
    missing = [
        key for key in ("ANTHROPIC_API_KEY", "SERPER_API_KEY")
        if not os.getenv(key)
    ]
    if missing:
        raise SystemExit(
            f"Missing environment variable(s): {', '.join(missing)}. "
            "Add them to your .env file."
        )

    postings = find_jobs(DEMO_ROLE, DEMO_LOCATION, DEMO_REMOTE, DEMO_HISTORY_DIR)
    print(f"\n\n=== FOUND {len(postings)} ROLE(S) ===\n")
    for job in postings:
        marker = " [NEW]" if job["is_new"] else ""
        salary = f" - {job['salary']}" if job["salary"] else ""
        print(f"- {job['title']} at {job['company']} ({job['location']}){marker}{salary}")
        print(f"  {job['link']}")
