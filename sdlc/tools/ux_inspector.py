"""
Playwright-backed tool that lets the ux_reviewer agent actually drive
the running Job Finder app, instead of only reasoning over text - see
the "Not yet wired up" note in sdlc/SDLC.py this closes for ux_reviewer
specifically (code_reviewer/local_tester/prod_tester/rollback_agent
still have no tools).

Covers the two real front-end "pages" of this single-page Streamlit
app (session_state["view"] == "main" or "format") - never the hidden
admin report page (?admin=1), which isn't part of the user-facing UX.

Every call is self-contained: it logs into (registering if needed) a
throwaway test account, ensures a fixture resume is on file, captures
the requested view, and then deletes that test account/resume before
returning - so no leftover test data survives a run, matching the
cleanup convention already required of local_tester and prod_tester.
"""

import json
import re
import sqlite3
import time
from pathlib import Path
from typing import Literal

from crewai.tools import BaseTool
from docx import Document
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
from playwright.sync_api import sync_playwright

# Generous, since these wait on a real Streamlit rerun (server round-trip
# + re-render), not just a fixed client-side delay - a loaded machine
# running several headless browsers at once can genuinely take a few
# seconds here, and waiting on the actual resulting state is more
# reliable than guessing a fixed sleep length.
STATE_CHANGE_TIMEOUT_MS = 15000

REPO_ROOT = Path(__file__).parent.parent.parent
DB_PATH = REPO_ROOT / "data" / "auth.db"
USERS_DIR = REPO_ROOT / "users"
SCREENSHOTS_DIR = Path(__file__).parent.parent / "ux_reports" / "screenshots"

TEST_EMAIL = "ux-reviewer-agent@example.com"
TEST_PASSWORD = "UxReviewerAgent123"
TEST_FIRST_NAME = "UX"
TEST_LAST_NAME = "ReviewerAgent"

# Elements whose computed style ties directly to a guideline in
# ux_guidelines.md (brand gradient, button styling, alert styling).
STYLE_TARGETS = {
    "hero_title": (".hero-title", ["color", "backgroundImage", "fontSize", "fontWeight"]),
    "hero_badge": (".hero-badge", ["color", "backgroundColor", "borderRadius"]),
    "primary_button": ("div.stButton button, div.stDownloadButton button",
                        ["backgroundColor", "backgroundImage", "color", "borderRadius", "boxShadow"]),
    "alert": ('[data-testid="stAlert"]', ["backgroundColor", "color"]),
}

PAGE_BACKGROUND_RGB = (10, 10, 15)  # #0a0a0f, from .streamlit/config.toml


def _rgb_from_css(value: str) -> tuple[int, int, int] | None:
    match = re.search(r"rgba?\(\s*(\d+),\s*(\d+),\s*(\d+)", value or "")
    if not match:
        return None
    return tuple(int(x) for x in match.groups())


def _relative_luminance(rgb: tuple[int, int, int]) -> float:
    def channel(c: int) -> float:
        c_norm = c / 255
        return c_norm / 12.92 if c_norm <= 0.03928 else ((c_norm + 0.055) / 1.055) ** 2.4

    r, g, b = rgb
    return 0.2126 * channel(r) + 0.7152 * channel(g) + 0.0722 * channel(b)


def _contrast_ratio(rgb_a: tuple[int, int, int], rgb_b: tuple[int, int, int]) -> float:
    l_a = _relative_luminance(rgb_a) + 0.05
    l_b = _relative_luminance(rgb_b) + 0.05
    return round(max(l_a, l_b) / min(l_a, l_b), 2)


def _make_fixture_resume(path: Path) -> None:
    doc = Document()
    doc.add_heading("UX Reviewer Test Resume", level=1)
    doc.add_paragraph("Experience: Senior Widget Wrangler at Testco (2020-2026).")
    doc.add_paragraph("Skills: automated UX review, being a disposable fixture.")
    doc.save(path)


def _delete_test_account_and_data() -> None:
    """Removes the throwaway test account row and its user directory,
    so no test data survives a review run."""
    if DB_PATH.exists():
        conn = sqlite3.connect(DB_PATH)
        try:
            conn.execute("DELETE FROM users WHERE username = ?", (TEST_EMAIL,))
            conn.commit()
        finally:
            conn.close()

    user_dir = USERS_DIR / TEST_EMAIL
    if user_dir.exists():
        for child in sorted(user_dir.rglob("*"), reverse=True):
            if child.is_file():
                child.unlink()
            else:
                child.rmdir()
        user_dir.rmdir()


def _wait_for_text(page, text: str, timeout: int = STATE_CHANGE_TIMEOUT_MS) -> bool:
    """Waits for text to actually appear (a real Streamlit rerun having
    happened), rather than guessing how long that rerun takes. Returns
    False instead of raising if it never shows up, so callers can
    branch on the outcome."""
    try:
        page.get_by_text(text, exact=False).first.wait_for(state="visible", timeout=timeout)
        return True
    except PlaywrightTimeoutError:
        return False


def _login_or_register(page) -> None:
    page.get_by_label("Email").fill(TEST_EMAIL)
    page.get_by_label("Password", exact=True).fill(TEST_PASSWORD)
    page.get_by_role("button", name="Login").click()

    if _wait_for_text(page, "Signed in as"):
        return

    # Not registered yet (or login failed) - switch to Register and
    # create the account.
    page.get_by_text("Register", exact=True).click()
    page.get_by_label("First name").wait_for(state="visible", timeout=STATE_CHANGE_TIMEOUT_MS)
    page.get_by_label("First name").fill(TEST_FIRST_NAME)
    page.get_by_label("Last name").fill(TEST_LAST_NAME)
    page.get_by_label("Email").fill(TEST_EMAIL)
    page.get_by_label("Password", exact=True).fill(TEST_PASSWORD)
    page.get_by_label("Repeat password").fill(TEST_PASSWORD)
    page.get_by_role("button", name="Register").click()
    if not _wait_for_text(page, "Registration successful"):
        raise RuntimeError("Registration did not confirm success within the timeout.")

    page.get_by_text("Login", exact=True).click()
    page.get_by_label("Email").wait_for(state="visible", timeout=STATE_CHANGE_TIMEOUT_MS)
    page.get_by_label("Email").fill(TEST_EMAIL)
    page.get_by_label("Password", exact=True).fill(TEST_PASSWORD)
    page.get_by_role("button", name="Login").click()
    if not _wait_for_text(page, "Signed in as"):
        raise RuntimeError("Login did not succeed within the timeout after registering.")


def _ensure_resume_uploaded(page, fixture_resume: Path) -> None:
    if page.locator("input[type='file']").count() > 0 and "Upload your resume" in page.inner_text("body"):
        page.locator("input[type='file']").first.set_input_files(str(fixture_resume))
        if not _wait_for_text(page, "Resume on file"):
            raise RuntimeError("Resume upload did not confirm within the timeout.")


def _capture_styles(page) -> dict:
    captured = {}
    for name, (selector, props) in STYLE_TARGETS.items():
        locator = page.locator(selector).first
        if locator.count() == 0:
            captured[name] = None
            continue
        values = page.evaluate(
            """([el, props]) => {
                const cs = getComputedStyle(el);
                const out = {};
                for (const p of props) out[p] = cs[p];
                return out;
            }""",
            [locator.element_handle(), props],
        )
        rgb = _rgb_from_css(values.get("color", ""))
        if rgb:
            values["contrast_vs_page_background"] = _contrast_ratio(rgb, PAGE_BACKGROUND_RGB)
        captured[name] = values
    return captured


class UXPageInspectorTool(BaseTool):
    name: str = "inspect_job_finder_page"
    description: str = (
        "Drives the real, running Job Finder app in a headless browser and "
        "returns what was actually observed on one of its two front-end "
        "views (never the admin page): 'main' (upload/review/search flow) "
        "or 'format' (resume format picker). Logs in with a throwaway test "
        "account (registering it if needed), uploads a fixture resume if "
        "one isn't already on file, navigates to the requested view, and "
        "returns a JSON string with: the visible text on that view, a "
        "screenshot file path, and computed CSS values (plus a measured "
        "WCAG contrast ratio against the page background) for the "
        "elements that this app's UX guidelines call out - the hero "
        "title/badge gradient, primary buttons, and any alert box. All "
        "test account/data created during the call is deleted before this "
        "tool returns, so nothing persists between or after calls."
    )

    def _run(self, view: Literal["main", "format"], base_url: str = "http://localhost:8502") -> str:
        SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)
        fixture_resume = SCREENSHOTS_DIR.parent / "fixture_resume.docx"
        if not fixture_resume.exists():
            _make_fixture_resume(fixture_resume)

        result = {"view": view, "base_url": base_url}
        try:
            with sync_playwright() as p:
                browser = p.chromium.launch()
                page = browser.new_page(viewport={"width": 1280, "height": 900})
                try:
                    page.goto(base_url, wait_until="networkidle", timeout=30000)
                    page.get_by_label("Email").wait_for(state="visible", timeout=STATE_CHANGE_TIMEOUT_MS)

                    _login_or_register(page)
                    _ensure_resume_uploaded(page, fixture_resume)

                    if view == "format":
                        page.get_by_role("button", name="Change resume format").click(
                            timeout=STATE_CHANGE_TIMEOUT_MS
                        )
                        if not _wait_for_text(page, "Pick a layout to rebuild"):
                            raise RuntimeError(
                                "Clicked 'Change resume format' but the format "
                                "view never rendered within the timeout."
                            )

                    screenshot_path = SCREENSHOTS_DIR / f"{view}_{int(time.time())}.png"
                    page.screenshot(path=str(screenshot_path), full_page=True)

                    result["screenshot_path"] = str(screenshot_path)
                    result["visible_text"] = page.inner_text("body")[:4000]
                    result["computed_styles"] = _capture_styles(page)
                finally:
                    browser.close()
        except Exception as e:
            result["error"] = f"{type(e).__name__}: {e}"
        finally:
            _delete_test_account_and_data()

        return json.dumps(result, indent=2)
